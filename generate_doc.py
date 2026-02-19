from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# --- COULEURS DU SITE ---
GOLD = RGBColor(0xD4, 0xAF, 0x37)
GOLD_DARK = RGBColor(0x8B, 0x69, 0x14)
DARK_BG = RGBColor(0x03, 0x00, 0x14)
DARK_CARD = RGBColor(0x0A, 0x06, 0x2A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
WHITE_DIM = RGBColor(0xAA, 0xAA, 0xBB)
ELECTRIC = RGBColor(0x00, 0xD4, 0xFF)
EMERALD = RGBColor(0x00, 0xC8, 0x96)
COSMIC = RGBColor(0x7B, 0x2F, 0xBE)

# --- PAGE BACKGROUND ---
def set_page_bg(doc, color_hex):
    """Set page background color"""
    bg = parse_xml(f'<w:background {nsdecls("w")} w:color="{color_hex}"/>')
    doc.element.insert(0, bg)

def set_cell_bg(cell, color_hex):
    """Set table cell background"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = GOLD
        run.font.name = 'Georgia'
    return h

def add_body_text(doc, text, bold=False, color=None, size=Pt(11), align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = size
    run.font.color.rgb = color or RGBColor(0x33, 0x33, 0x33)
    run.bold = bold
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet(doc, text, color=None):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.color.rgb = color or RGBColor(0x44, 0x44, 0x44)
    return p

def make_dark_table(doc, headers, rows, accent=GOLD):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, '030014')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h_text)
        run.font.color.rgb = GOLD
        run.font.name = 'Georgia'
        run.font.size = Pt(10)
        run.bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.add_row().cells
        bg = '0A062A' if row_idx % 2 == 0 else '0F0B35'
        for i, val in enumerate(row_data):
            cell = row_cells[i]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.color.rgb = WHITE_DIM
            run.font.name = 'Calibri'
            run.font.size = Pt(9.5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="D4AF37"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D4AF37"/>'
        '  <w:insideH w:val="single" w:sz="2" w:space="0" w:color="1A1640"/>'
        '  <w:insideV w:val="single" w:sz="2" w:space="0" w:color="1A1640"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    return table

def add_section_divider(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\u2500' * 40)
    run.font.color.rgb = RGBColor(0xD4, 0xAF, 0x37)
    run.font.size = Pt(8)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

# --- Margins ---
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ============================================================
#                      PAGE DE GARDE
# ============================================================
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\u2316')
run.font.size = Pt(48)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('TIMETRAVEL AGENCY')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = GOLD
run.font.name = 'Georgia'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\u2500' * 30)
run.font.color.rgb = GOLD
run.font.size = Pt(10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Webapp Interactive')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = 'Georgia'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Projet Final \u2014 M1/M2 Digital & IA')
run.font.size = Pt(13)
run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('R\u00e9alis\u00e9 par')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('YAHIA Rayan')
run.font.size = Pt(18)
run.bold = True
run.font.color.rgb = GOLD_DARK
run.font.name = 'Georgia'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\u2192  https://timetravel-agency-ten.vercel.app')
run.font.size = Pt(11)
run.font.color.rgb = ELECTRIC
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('19 f\u00e9vrier 2026')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = 'Calibri'

doc.add_page_break()

# ============================================================
#                        SOMMAIRE
# ============================================================
add_styled_heading(doc, '\u2316  Sommaire', level=1)
add_section_divider(doc)

sommaire = [
    ('1.', 'Pr\u00e9sentation du projet'),
    ('2.', 'Stack technique'),
    ('3.', 'Architecture & Planning \u2014 Phase 1'),
    ('4.', 'G\u00e9n\u00e9ration de code & Vibe Coding \u2014 Phase 2'),
    ('5.', 'Intelligence Artificielle & Agents \u2014 Phase 3'),
    ('6.', 'Features impl\u00e9ment\u00e9es'),
    ('7.', 'D\u00e9ploiement \u2014 Phase 4'),
    ('8.', 'Outils IA utilis\u00e9s'),
    ('9.', 'Conclusion'),
]
for num, title in sommaire:
    p = doc.add_paragraph()
    run = p.add_run(num + '  ')
    run.font.color.rgb = GOLD
    run.font.name = 'Georgia'
    run.font.size = Pt(12)
    run.bold = True
    run = p.add_run(title)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(6)

doc.add_page_break()

# ============================================================
#                1. PRESENTATION
# ============================================================
add_styled_heading(doc, '\u25C8  1. Pr\u00e9sentation du projet', level=1)
add_section_divider(doc)

add_body_text(doc,
    'TimeTravel Agency est une webapp interactive moderne pr\u00e9sentant une agence fictive de voyages temporels de luxe. '
    'Le site permet aux visiteurs de d\u00e9couvrir trois destinations temporelles uniques, '
    "d'interagir avec un agent conversationnel intelligent, de personnaliser leur voyage via un quiz, "
    'et de simuler une r\u00e9servation compl\u00e8te.'
)

add_body_text(doc,
    "Le projet a \u00e9t\u00e9 d\u00e9velopp\u00e9 dans le cadre du cours Digital & IA, en combinant des techniques de vibe coding, "
    "d'intelligence artificielle g\u00e9n\u00e9rative, et de d\u00e9veloppement front-end moderne."
)

# Destinations highlight
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\u2605 Paris 1889')
run.font.color.rgb = GOLD
run.font.size = Pt(13)
run.bold = True
run.font.name = 'Georgia'
run = p.add_run('   \u00b7   ')
run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
run = p.add_run('\u2605 Cr\u00e9tac\u00e9')
run.font.color.rgb = EMERALD
run.font.size = Pt(13)
run.bold = True
run.font.name = 'Georgia'
run = p.add_run('   \u00b7   ')
run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
run = p.add_run('\u2605 Florence 1504')
run.font.color.rgb = RGBColor(0xC4, 0x1E, 0x3A)
run.font.size = Pt(13)
run.bold = True
run.font.name = 'Georgia'

doc.add_page_break()

# ============================================================
#                2. STACK TECHNIQUE
# ============================================================
add_styled_heading(doc, '\u2699  2. Stack technique', level=1)
add_section_divider(doc)

make_dark_table(doc, ['Technologie', 'Usage'], [
    ('React 19', 'Framework UI \u2014 composants, state management'),
    ('TypeScript', 'Typage statique, fiabilit\u00e9 du code'),
    ('Tailwind CSS v4', 'Styling utility-first, design responsive'),
    ('Framer Motion', 'Animations fluides, transitions, micro-interactions'),
    ('Lucide React', 'Ic\u00f4nes SVG coh\u00e9rentes'),
    ('Vite 7', 'Build tool ultra-rapide, HMR'),
    ('Web Audio API', "Son d'ambiance g\u00e9n\u00e9ratif"),
    ('Vercel', 'D\u00e9ploiement et h\u00e9bergement production'),
    ('Claude Code', 'Assistant IA (Claude Opus 4.6)'),
])

doc.add_page_break()

# ============================================================
#                3. PHASE 1
# ============================================================
add_styled_heading(doc, '\u25B6  3. Architecture & Planning \u2014 Phase 1', level=1)
add_section_divider(doc)

add_styled_heading(doc, '3.1  D\u00e9finition des features', level=2)
add_body_text(doc, 'Fonctionnalit\u00e9s d\u00e9finies pour la webapp :')

features_list = [
    "\u25B8  Hero section avec animation typewriter et particules",
    "\u25B8  \u00c0 propos avec compteurs anim\u00e9s (3 \u00e9poques, 2 847 voyageurs, 100% retours)",
    "\u25B8  3 destinations avec cards interactives et compteur de places",
    "\u25B8  Galerie immersive avec onglets filtres (12 cartes)",
    "\u25B8  Frise chronologique interactive des 3 \u00e9poques",
    "\u25B8  Quiz de recommandation personnalis\u00e9e (4 questions)",
    "\u25B8  Formulaire de r\u00e9servation multi-\u00e9tapes",
    "\u25B8  Chatbot IA conversationnel (Chronos)",
    "\u25B8  Section t\u00e9moignages (6 avis de voyageurs)",
    "\u25B8  FAQ en accord\u00e9on anim\u00e9 (8 questions)",
    "\u25B8  Mode sombre/clair avec toggle",
    "\u25B8  Son d'ambiance g\u00e9n\u00e9ratif (Web Audio API)",
    '\u25B8  Easter egg clavier (taper "time")',
    "\u25B8  Page 404 th\u00e9matique",
]
for f in features_list:
    add_bullet(doc, f)

add_styled_heading(doc, '3.2  Structure de navigation', level=2)
add_body_text(doc, 'Le site suit un flow vertical single-page responsive (mobile-first) :')

nav_items = [
    'Header fixe \u2192 Hero \u2192 \u00c0 propos \u2192 Destinations \u2192 Galerie',
    'Timeline \u2192 Quiz \u2192 R\u00e9servation \u2192 T\u00e9moignages \u2192 FAQ \u2192 Footer',
    'Overlay : Chatbot (bas droite) + Son ambiance (bas gauche) + Easter egg',
]
for n in nav_items:
    add_bullet(doc, n)

doc.add_page_break()

# ============================================================
#                4. PHASE 2
# ============================================================
add_styled_heading(doc, '\u25C6  4. G\u00e9n\u00e9ration de code & Vibe Coding \u2014 Phase 2', level=1)
add_section_divider(doc)

add_styled_heading(doc, '4.1  Setup & G\u00e9n\u00e9ration initiale', level=2)
add_body_text(doc,
    "Le projet a \u00e9t\u00e9 initialis\u00e9 avec Vite + React + TypeScript, puis d\u00e9velopp\u00e9 "
    "it\u00e9rativement avec Claude Code (Opus 4.6). Chaque composant a \u00e9t\u00e9 g\u00e9n\u00e9r\u00e9 "
    "via des prompts d\u00e9taill\u00e9s, test\u00e9, puis int\u00e9gr\u00e9 dans l'application."
)

add_styled_heading(doc, '4.2  Int\u00e9gration des assets', level=2)
add_body_text(doc,
    "Les visuels sont g\u00e9n\u00e9r\u00e9s en CSS pur (gradients, ombres, glow effects) avec des emojis "
    "comme \u00e9l\u00e9ments visuels. Chaque \u00e9poque a sa palette de couleurs : "
    "gold pour Paris, emerald pour le Cr\u00e9tac\u00e9, rouge renaissance pour Florence."
)

add_styled_heading(doc, '4.3  Animations (exercice optionnel \u2014 r\u00e9alis\u00e9 \u2713)', level=2)
add_body_text(doc, 'Animations impl\u00e9ment\u00e9es avec Framer Motion :')

anims = [
    'Fade-in progressif au scroll (whileInView)',
    'Typewriter anim\u00e9 sur le sous-titre Hero',
    'Hover effects : scale, glow, fl\u00e8che directionnelle',
    'Transitions entre \u00e9tapes (quiz, r\u00e9servation)',
    'Compteurs anim\u00e9s, accord\u00e9on FAQ',
    'Particules flottantes en arri\u00e8re-plan',
    'Easter egg : effet tunnel warp temporel',
    'Loading screen avec animation d\u2019entr\u00e9e',
]
for a in anims:
    add_bullet(doc, a)

doc.add_page_break()

# ============================================================
#                5. PHASE 3
# ============================================================
add_styled_heading(doc, '\u25C8  5. Intelligence Artificielle & Agents \u2014 Phase 3', level=1)
add_section_divider(doc)

add_styled_heading(doc, '5.1  Agent conversationnel \u2014 Chronos', level=2)

add_body_text(doc,
    "Le chatbot Chronos est accessible via un widget flottant en bas \u00e0 droite. "
    "Il utilise un syst\u00e8me de pattern matching pour r\u00e9pondre intelligemment aux visiteurs."
)

add_body_text(doc, 'Capacit\u00e9s de Chronos :', bold=True)
chat_features = [
    'R\u00e9ponses d\u00e9taill\u00e9es sur les 3 destinations',
    'Informations tarifs (12 500 \u20ac / 18 900 \u20ac / 14 200 \u20ac)',
    'S\u00e9curit\u00e9, garanties, bagages, dur\u00e9e',
    'Recommandation selon les int\u00e9r\u00eats',
    'Quick actions + indicateur de frappe',
]
for c in chat_features:
    add_bullet(doc, c)

add_body_text(doc,
    "Personnalit\u00e9 : professionnel, chaleureux, passionn\u00e9 d\u2019histoire, "
    "expert en voyage temporel. Ton enthousiaste sans \u00eatre familier.",
    color=RGBColor(0x66, 0x66, 0x66), size=Pt(10)
)

add_styled_heading(doc, '5.2  Quiz de recommandation (exercice optionnel \u2014 r\u00e9alis\u00e9 \u2713)', level=2)
add_body_text(doc,
    "Quiz de 4 questions avec syst\u00e8me de scoring. "
    "Chaque r\u00e9ponse attribue des points aux 3 destinations. "
    "La destination avec le meilleur score est recommand\u00e9e avec une explication personnalis\u00e9e."
)

make_dark_table(doc, ['Question', 'Options'], [
    ("Type d'exp\u00e9rience", 'Culturelle / Aventure / \u00c9l\u00e9gance'),
    ('P\u00e9riode pr\u00e9f\u00e9r\u00e9e', 'Moderne / Ancienne / Renaissance'),
    ('Pr\u00e9f\u00e9rence', 'Urbain / Nature / Art'),
    ('Activit\u00e9 id\u00e9ale', 'Monuments / Faune / Mus\u00e9es'),
])

doc.add_page_break()

# ============================================================
#                6. FEATURES
# ============================================================
add_styled_heading(doc, '\u2726  6. Features impl\u00e9ment\u00e9es', level=1)
add_section_divider(doc)

make_dark_table(doc, ['Feature', 'Description', 'Fichier'], [
    ('Loading Screen', '\u00c9cran de chargement anim\u00e9', 'LoadingScreen.tsx'),
    ('Particules', 'Arri\u00e8re-plan anim\u00e9', 'ParticleBackground.tsx'),
    ('Header', 'Nav fixe + scroll spy + th\u00e8me', 'Header.tsx'),
    ('Hero', 'Titre anim\u00e9, typewriter, CTA', 'Hero.tsx'),
    ('\u00c0 propos', '3 compteurs anim\u00e9s', 'About.tsx'),
    ('Destinations', '3 cards + prix + places', 'Destinations.tsx'),
    ('Galerie', '12 cartes, onglets par \u00e9poque', 'Gallery.tsx'),
    ('Timeline', 'Frise chronologique verticale', 'Timeline.tsx'),
    ('Quiz', 'Recommandation 4 questions', 'Quiz.tsx'),
    ('R\u00e9servation', 'Formulaire 4 \u00e9tapes', 'Booking.tsx'),
    ('T\u00e9moignages', '6 avis + \u00e9toiles', 'Testimonials.tsx'),
    ('FAQ', '8 questions accord\u00e9on', 'FAQ.tsx'),
    ('Chatbot', 'Agent Chronos', 'Chatbot.tsx'),
    ('Th\u00e8me', 'Toggle sombre/clair', 'ThemeToggle.tsx'),
    ('Son ambiance', 'Web Audio API', 'AmbientSound.tsx'),
    ('Easter egg', 'Taper "time"', 'EasterEgg.tsx'),
    ('Page 404', 'Perdu dans le temps', 'NotFound.tsx'),
])

doc.add_page_break()

# ============================================================
#                7. DEPLOIEMENT
# ============================================================
add_styled_heading(doc, '\u2192  7. D\u00e9ploiement \u2014 Phase 4', level=1)
add_section_divider(doc)

add_body_text(doc, 'Processus de d\u00e9ploiement :')
deploy = [
    'Build de production : npm run build (TypeScript + Vite)',
    'D\u00e9ploiement via Vercel CLI : vercel --prod',
    'CDN mondial pour des temps de chargement optimaux',
    'HTTPS automatique + certificat SSL',
]
for d in deploy:
    add_bullet(doc, d)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('URL de production')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('https://timetravel-agency-ten.vercel.app')
run.font.size = Pt(16)
run.font.color.rgb = GOLD
run.font.name = 'Georgia'
run.bold = True

doc.add_page_break()

# ============================================================
#                8. OUTILS IA
# ============================================================
add_styled_heading(doc, '\u2699  8. Outils IA utilis\u00e9s', level=1)
add_section_divider(doc)

add_body_text(doc, 'Transparence sur les outils IA utilis\u00e9s dans ce projet :')

make_dark_table(doc, ['Outil', 'Mod\u00e8le', 'Usage'], [
    ('Claude Code', 'Claude Opus 4.6', 'G\u00e9n\u00e9ration du code, debugging, d\u00e9ploiement'),
    ('Framer Motion', 'Open source', 'Animations et transitions'),
    ('Tailwind CSS v4', 'Open source', 'Framework CSS utility-first'),
    ('Vercel', 'Cloud', 'H\u00e9bergement et CDN'),
    ('Web Audio API', 'Native', 'Son g\u00e9n\u00e9ratif navigateur'),
])

doc.add_page_break()

# ============================================================
#                9. CONCLUSION
# ============================================================
add_styled_heading(doc, '\u25C8  9. Conclusion', level=1)
add_section_divider(doc)

add_body_text(doc,
    "Ce projet a permis d\u2019explorer le vibe coding et l\u2019utilisation d\u2019outils IA pour le d\u00e9veloppement web. "
    "La webapp TimeTravel Agency int\u00e8gre toutes les fonctionnalit\u00e9s demand\u00e9es dans le brief : "
    "interface immersive avec animations, agent conversationnel, quiz de recommandation, "
    "formulaire de r\u00e9servation, et de nombreux bonus."
)

add_body_text(doc,
    "Le site est enti\u00e8rement responsive, d\u00e9ploy\u00e9 en production sur Vercel, "
    "et propose une exp\u00e9rience utilisateur premium avec 17 composants interactifs."
)

doc.add_paragraph()
add_section_divider(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('"Le pass\u00e9 n\u2019attend que vous."')
run.font.size = Pt(14)
run.font.color.rgb = GOLD
run.font.name = 'Georgia'
run.italic = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\u2014 TimeTravel Agency')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = 'Georgia'

# --- SAVE ---
doc.save(r'C:\Users\Rayan\Documents\TimeTravel-Agency\TimeTravel_Agency_Rendu.docx')
print('Document genere avec succes !')
